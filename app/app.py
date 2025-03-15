from flask import Flask, render_template, request, redirect, send_file, url_for, flash, jsonify
from models import db, Exam, User
from collections import defaultdict
import json
from docx import Document
from docx2pdf import convert
import os
import pythoncom 

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///exams.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.secret_key = 'sua_chave_secreta_aqui'


db.init_app(app)


current_dir = os.path.dirname(os.path.abspath(__file__))
categories_path = os.path.join(current_dir, 'static', 'infos', 'categorias.json')
subcategories_path = os.path.join(current_dir, 'static', 'infos', 'subcategories.json')


with open(categories_path, 'r', encoding='utf-8') as file:
    EXAM_TYPES = json.load(file)

with open(subcategories_path, 'r', encoding='utf-8') as file:
    EXAM_SUBCATEGORIES = json.load(file)


@app.route('/')
def root():
    return redirect(url_for('home'))


@app.route('/home')
def home():
    return render_template('index.html')


@app.route('/exame')
def exame():
    return render_template('exame.html', exam_types=EXAM_TYPES, exam_subcategories=EXAM_SUBCATEGORIES)


@app.route('/generate_label', methods=['GET', 'POST'])
def generate_label():
    if request.method == 'POST':
        patient_name = request.form['patient_name']
        exam_type = request.form['exam_type']
        return render_template('label.html', patient_name=patient_name, exam_type=exam_type)
    return render_template('generate_label.html', exam_types=EXAM_TYPES)


@app.route('/get_patient_details')
def get_patient_details():
    patient_name = request.args.get('patient_name')
    patient = Exam.query.filter_by(patient_name=patient_name).first()

    if patient:
        return jsonify({
            'cpf': patient.cpf,
            'phone': patient.phone,
            'address': patient.address
        })
    else:
        return jsonify({}), 404



@app.route('/edit/<int:exam_id>', methods=['GET', 'POST'])
def edit_exam(exam_id):
    exam = Exam.query.get_or_404(exam_id)  

    if request.method == 'POST':
        
        exam.patient_name = request.form['patient_name']
        exam.category = request.form['category']
        exam.subcategory = request.form['subcategory']
        exam.result = request.form['result']

        
        details = {}
        for key, value in request.form.items():
            if key.startswith('detail_key_') and value:
                detail_number = key.split('_')[-1]
                detail_value = request.form.get(f'detail_value_{detail_number}', '')
                if detail_value:
                    details[value] = detail_value
        exam.details = details

        db.session.commit()  
        return redirect(url_for('results'))

    
    return render_template('edit_exam.html', exam=exam, exam_types=EXAM_TYPES, exam_subcategories=EXAM_SUBCATEGORIES)





@app.route('/submit', methods=['POST'])
def submit_exam():
    
    patient_name = request.form['patient_name']
    category = request.form['category']
    subcategory = request.form['subcategory']
    result = request.form.get('result', '')  

    
    user = User.query.filter_by(patient_name=patient_name).first()

    if not user:
        flash('Paciente não encontrado. Por favor, cadastre o paciente primeiro.', 'error')
        return redirect(url_for('exame'))

    
    details = {}
    for key, value in request.form.items():
        if key.startswith('detail_key_') and value:  
            detail_number = key.split('_')[-1]  
            detail_value = request.form.get(f'detail_value_{detail_number}', '')  
            if detail_value:  
                details[value] = detail_value

    
    new_exam = Exam(
        patient_name=patient_name,  
        category=category,
        subcategory=subcategory,
        result=result,  
        details=details  
    )

    db.session.add(new_exam)
    db.session.commit()
    return redirect(url_for('results'))


@app.route('/results')
def results():
    exams = Exam.query.all()
    grouped_exams = defaultdict(lambda: {"category": None, "exams": [], "patient_info": {}})
    
    for exam in exams:
        
        user = User.query.filter_by(patient_name=exam.patient_name).first()
        
        if user:
            grouped_exams[exam.patient_name]["patient_info"] = {
                "cpf": user.cpf,
                "phone": user.phone,
                "address": user.address
            }
        
        if not grouped_exams[exam.patient_name]["category"]:
            grouped_exams[exam.patient_name]["category"] = exam.category
        
        grouped_exams[exam.patient_name]["exams"].append({
            "id": exam.id,
            "subcategory": exam.subcategory,
            "result": exam.result,
            "details": exam.details
        })
    
    return render_template('result.html', grouped_exams=grouped_exams)


@app.route('/client_record_search', methods=['GET', 'POST'])
def client_record_search():
    if request.method == 'POST':
        patient_name = request.form['patient_name']
        cpf = request.form['cpf']
        phone = request.form['phone']
        address = request.form['address']
        category = request.form['category']
        subcategory = request.form['subcategory']
        
        
        new_exam = User(
            patient_name=patient_name,
            cpf=cpf,
            phone=phone,
            address=address,
            category=category,
            subcategory=subcategory
        )
        db.session.add(new_exam)
        db.session.commit()
        
        flash('Ficha do cliente e exame salvos com sucesso!', 'success')
        return redirect(url_for('client_record', patient_name=patient_name))
    
    return render_template('client_record_search.html', exam_types=EXAM_TYPES, exam_subcategories=EXAM_SUBCATEGORIES)

@app.route('/delete/<int:exam_id>', methods=['POST'])
def delete_exam(exam_id):
    exam = Exam.query.get_or_404(exam_id)  
    db.session.delete(exam)  
    db.session.commit()  
    return redirect(url_for('results'))


@app.route('/client_record/<string:patient_name>')
def client_record(patient_name):
    exams = User.query.filter_by(patient_name=patient_name).all()
    
    if not exams:
        flash(f'Nenhum exame encontrado para o paciente: {patient_name}', 'warning')
        return redirect(url_for('client_record_search'))
    
    
    client_info = {
        "cpf": exams[0].cpf,
        "phone": exams[0].phone,
        "address": exams[0].address
    }
    
    
    grouped_exams = defaultdict(lambda: {"exams": []})
    for exam in exams:
        grouped_exams[exam.category]["exams"].append({
            "id": exam.id,
            "subcategory": exam.subcategory,
        })
    
    return render_template(
        'client_record.html',
        patient_name=patient_name,
        cpf=client_info["cpf"],
        phone=client_info["phone"],
        address=client_info["address"],
        grouped_exams=grouped_exams
    )

from docx import Document
from docx2pdf import convert
import os
import pythoncom  # Importe o pythoncom para inicializar o CoInitialize

@app.route('/generate_word/<string:patient_name>')
def generate_word(patient_name):
    # Buscar as informações do paciente
    user = User.query.filter_by(patient_name=patient_name).first()
    exams = Exam.query.filter_by(patient_name=patient_name).all()

    if not user or not exams:
        flash('Paciente ou exames não encontrados.', 'error')
        return redirect(url_for('results'))

    # Criar um documento Word
    doc = Document()
    doc.add_heading(f'Relatório de Exames - {patient_name}', 0)

    # Adicionar informações do paciente
    doc.add_heading('Informações do Paciente', level=1)
    doc.add_paragraph(f'Nome: {user.patient_name}')
    doc.add_paragraph(f'CPF: {user.cpf}')
    doc.add_paragraph(f'Telefone: {user.phone}')
    doc.add_paragraph(f'Endereço: {user.address}')

    # Adicionar informações dos exames
    doc.add_heading('Exames Realizados', level=1)
    for exam in exams:
        doc.add_heading(f'Categoria: {exam.category}', level=2)
        doc.add_paragraph(f'Subcategoria: {exam.subcategory}')
        doc.add_paragraph(f'Resultado: {exam.result}')
        if exam.details:
            doc.add_paragraph('Detalhes:')
            for key, value in exam.details.items():
                doc.add_paragraph(f'{key}: {value}')

    # Salvar o documento Word
    current_dir = os.path.dirname(os.path.abspath(__file__))
    word_path = os.path.join(current_dir, 'static', 'reports', f'{patient_name.replace(" ", "_")}_report.docx')
    pdf_path = os.path.join(current_dir, 'static', 'reports', f'{patient_name.replace(" ", "_")}_report.pdf')
    
    # Certifique-se de que o diretório existe
    os.makedirs(os.path.dirname(word_path), exist_ok=True)
    
    doc.save(word_path)

    # Inicializar o CoInitialize antes de converter o Word para PDF
    pythoncom.CoInitialize()  # Inicializa o CoInitialize

    try:
        # Converter o Word para PDF
        convert(word_path, pdf_path)
    finally:
        pythoncom.CoUninitialize()  # Finaliza o CoInitialize

    # Retornar o arquivo PDF para download
    return send_file(pdf_path, as_attachment=True)

if __name__ == '__main__':
    with app.app_context():
        db.create_all()  
    app.run(debug=True)